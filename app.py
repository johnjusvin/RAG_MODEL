# Required Libraries
import streamlit as st
import pandas as pd
import os
import datetime
from sqlalchemy import create_engine, Column, Integer, String, Text, ForeignKey, DateTime
from sqlalchemy.orm import sessionmaker, declarative_base, relationship, Session
from sqlalchemy.exc import IntegrityError # For handling potential foreign key constraints

# For ChromaDB
import chromadb

# For document text extraction
from pypdf import PdfReader  # For PDFs
from docx import Document as DocxDocument  # To avoid conflict with SQLAlchemy Document model
import io  # To handle file-like objects for text extraction


# SQLAlchemy Database Setup
DATABASE_URL = "postgresql://postgres:21112005POST@localhost:5432/postgres" # Update with your PostgreSQL credentials
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)
Base = declarative_base()

# Database Models
class Knowledge(Base):
    __tablename__ = 'knowledge'
    id = Column(Integer, primary_key=True)
    name = Column(String, nullable=False)
    description = Column(Text)
    documents = relationship("Document", back_populates="knowledge", cascade="all, delete")

class Document(Base):
    __tablename__ = 'document'
    id = Column(Integer, primary_key=True)
    knowledge_id = Column(Integer, ForeignKey('knowledge.id'))
    name = Column(String)
    filetype = Column(String)
    size = Column(Integer)
    path = Column(String)
    uploaded_at = Column(DateTime, default=datetime.datetime.utcnow)
    knowledge = relationship("Knowledge", back_populates="documents")

# Create tables in PostgreSQL (run once - SQLAlchemy handles existence)
Base.metadata.create_all(bind=engine)


# ChromaDB Vector Database Setup
CHROMA_PERSIST_DIR = "chroma_db" # Directory to store ChromaDB data
chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)

# Get or create a collection for your documents in ChromaDB
# This single line handles both getting an existing collection or creating a new one
documents_collection = chroma_client.get_or_create_collection(name="knowledge_documents")


# Helper Function for Text Extraction from Files
def extract_text_from_file(file_path, filetype):
    """
    Extracts text content from various file types.
    Currently supports PDF, DOCX, and plain text.
    """
    text = ""
    try:
        if filetype == "application/pdf":
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or "" # Add empty string if None
        elif filetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document": # .docx
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif filetype == "text/plain":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            st.warning(f"Unsupported file type for text extraction: {filetype}. Skipping text indexing in ChromaDB.")
    except Exception as e:
        st.error(f"Error extracting text from {file_path}: {e}")
        text = "" # Ensure text is empty on error
    return text


# Initialize session state variables for upload tracking
if 'show_knowledge_form' not in st.session_state:
    st.session_state.show_knowledge_form = False
if 'upload_for_id' not in st.session_state:
    st.session_state.upload_for_id = None
if 'pending_vectorization_doc' not in st.session_state:
    st.session_state.pending_vectorization_doc = None # To store info of doc uploaded to Postgres, but not yet Chroma

# Sidebar Navigation
with st.sidebar:
    st.title("Navigation")
    page = st.radio("Go to", ["Knowledge", "Chat"])

    # This block is for the "Add Knowledge" form
    if st.session_state.get("show_knowledge_form", False):
        st.markdown("### Add Knowledge")
        name_input = st.text_input("Name")
        desc_input = st.text_area("Description")
        if st.button("Save Knowledge"):
            db: Session = SessionLocal()
            new_k = Knowledge(name=name_input, description=desc_input)
            db.add(new_k)
            db.commit()
            db.close()
            st.session_state.show_knowledge_form = False
            st.rerun()

    # This block is for the "Upload to:" file uploader (for PostgreSQL)
    if st.session_state.get("upload_for_id"):
        st.markdown(f"### Upload to: {st.session_state.upload_for_name}")
        st.markdown(f"_Description: {st.session_state.upload_for_desc}_")
        uploaded_file = st.file_uploader("Choose a file")
        if uploaded_file:
            # Define storage path based on knowledge name
            storage_dir = os.path.join("storage", st.session_state.upload_for_name.replace(" ", "_").lower()) # Sanitize name for path
            os.makedirs(storage_dir, exist_ok=True)
            file_path = os.path.join(storage_dir, uploaded_file.name)

            # Save the file to local storage
            try:
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                st.success(f"File '{uploaded_file.name}' saved to local storage.")
            except Exception as e:
                st.error(f"Error saving file to disk: {e}")
                st.session_state.upload_for_id = None # Reset state
                st.stop() # Stop execution if file cannot be saved

            # Insert document metadata into PostgreSQL
            db: Session = SessionLocal()
            try:
                new_doc = Document(
                    knowledge_id=st.session_state.upload_for_id,
                    name=uploaded_file.name,
                    filetype=uploaded_file.type,
                    size=uploaded_file.size,
                    path=file_path,
                    uploaded_at=datetime.datetime.now()
                )
                db.add(new_doc)
                db.commit()

                # Store details of the newly uploaded document for pending vectorization
                st.session_state.pending_vectorization_doc = {
                    "document_id": new_doc.id,
                    "knowledge_id": st.session_state.upload_for_id,
                    "knowledge_name": st.session_state.upload_for_name,
                    "knowledge_description": st.session_state.upload_for_desc,
                    "file_name": uploaded_file.name,
                    "file_type": uploaded_file.type,
                    "size": uploaded_file.size,
                    "path": file_path,
                    "uploaded_at": str(datetime.datetime.now()) # Stored as string for ChromaDB compatibility
                }
                st.success("Document metadata saved to PostgreSQL. Now, click 'Add to VectorDB' below.")

            except Exception as e:
                db.rollback() # Rollback changes if an error occurs during DB insert
                st.error(f"Error saving document metadata to PostgreSQL: {e}")
            finally:
                db.close() # Always close the session after this DB operation

            # Reset upload form state, but keep pending_vectorization_doc
            st.session_state.upload_for_id = None
            st.rerun() # Rerun to update the UI and show the new button


    # --- New section for "Add to VectorDB" button in the sidebar ---
    if st.session_state.get("pending_vectorization_doc"):
        doc_info = st.session_state.pending_vectorization_doc
        st.markdown("---") # Separator
        st.markdown("### Process for VectorDB")
        st.info(f"Ready to index: '{doc_info['file_name']}'")

        if st.button("Add to VectorDB"):
            extracted_text = extract_text_from_file(doc_info['path'], doc_info['file_type'])

            if extracted_text:
                try:
                    documents_collection.add(
                        documents=[extracted_text],
                        metadatas=[{
                            "knowledge_id": doc_info['knowledge_id'],
                            "knowledge_name": doc_info['knowledge_name'],
                            "knowledge_description": doc_info['knowledge_description'],
                            "document_id": doc_info['document_id'],
                            "file_name": doc_info['file_name'],
                            "file_type": doc_info['file_type'],
                            "size": doc_info['size'],
                            "path": doc_info['path'],
                            "uploaded_at": doc_info['uploaded_at'] # Use the string formatted datetime
                        }],
                        ids=[f"doc_{doc_info['document_id']}"] # Unique ID for ChromaDB
                    )
                    st.success(f"'{doc_info['file_name']}' successfully added to VectorDB (ChromaDB)!")
                    st.session_state.pending_vectorization_doc = None # Clear pending state
                    st.rerun() # Rerun to remove the button
                except Exception as e:
                    st.error(f"Error adding '{doc_info['file_name']}' to VectorDB: {e}")
            else:
                st.warning(f"Could not extract text from '{doc_info['file_name']}'. Cannot add to VectorDB.")


# Main Content Area
if page == "Knowledge":
    st.title("Knowledge")

    # Header with Add Button
    col1, col2 = st.columns([8, 1])
    with col1:
        st.subheader("Knowledge Table")
    with col2:
        if st.button("+"):
            st.session_state.show_knowledge_form = True

    # Fetch Knowledge Table
    db: Session = SessionLocal() # Open the session BEFORE the loop
    knowledge_list = db.query(Knowledge).order_by(Knowledge.id.desc()).all()

    # Display Knowledge Bases in Expanders with their Documents
    if not knowledge_list:
        st.info("No Knowledge Bases created yet. Click 'Add New' to get started!")

    for k in knowledge_list:
        with st.expander(f"**{k.name}** – {k.description}"):
            st.markdown("##### Associated Documents:")

            if k.documents:
                doc_data = []
                for doc in k.documents:
                    doc_data.append({
                        "id": doc.id, # Keep ID for selection
                        "File Name": doc.name,
                        "Type": doc.filetype,
                        "Size (bytes)": doc.size,
                        "Path": doc.path,
                        "Uploaded At": doc.uploaded_at.strftime("%Y-%m-%d %H:%M:%S")
                    })
                doc_df = pd.DataFrame(doc_data)
                st.dataframe(doc_df[['File Name', 'Type', 'Size (bytes)', 'Uploaded At']], use_container_width=True)

                # --- Document Removal Section ---
                st.markdown("---")
                st.markdown("##### Remove Document:")
                document_options = {f"{d.name} (ID: {d.id})": d.id for d in k.documents}
                selected_doc_key = st.selectbox(
                    f"Select document to remove from **{k.name}**:",
                    options=list(document_options.keys()),
                    key=f"remove_select_{k.id}" # Unique key for each selectbox
                )

                if selected_doc_key:
                    doc_to_remove_id = document_options[selected_doc_key]
                    if st.button(f"Remove Selected Document", key=f"remove_btn_{k.id}"):
                        doc_to_remove = db.query(Document).filter_by(id=doc_to_remove_id).first()

                        if doc_to_remove:
                            # 1. Delete from ChromaDB
                            try:
                                chroma_doc_id = f"doc_{doc_to_remove.id}"
                                documents_collection.delete(ids=[chroma_doc_id])
                                st.success(f"Document '{doc_to_remove.name}' removed from VectorDB (ChromaDB).")
                            except Exception as e:
                                st.warning(f"Could not remove '{doc_to_remove.name}' from VectorDB (ChromaDB). It might not have been indexed or an error occurred: {e}")

                            # 2. Delete file from local storage
                            if os.path.exists(doc_to_remove.path):
                                try:
                                    os.remove(doc_to_remove.path)
                                    st.success(f"File '{doc_to_remove.name}' deleted from local storage.")
                                except Exception as e:
                                    st.error(f"Error deleting file '{doc_to_remove.name}' from storage: {e}")
                            else:
                                st.info(f"File '{doc_to_remove.name}' not found in local storage path: {doc_to_remove.path}")


                            # 3. Delete from PostgreSQL
                            try:
                                db.delete(doc_to_remove)
                                db.commit()
                                st.success(f"Document '{doc_to_remove.name}' removed from PostgreSQL.")
                                st.rerun() # Rerun to update the displayed list
                            except IntegrityError as e:
                                db.rollback()
                                st.error(f"Integrity Error: Could not delete document from PostgreSQL. {e}")
                            except Exception as e:
                                db.rollback()
                                st.error(f"Error deleting document from PostgreSQL: {e}")
                        else:
                            st.error("Document not found in database.")
            else:
                st.info("No documents uploaded yet.")

            # Button to trigger document upload
            if st.button(f"Upload Document to **{k.name}**", key=f"upload_{k.id}"):
                st.session_state.upload_for_id = k.id
                st.session_state.upload_for_name = k.name
                st.session_state.upload_for_desc = k.description
                st.session_state.pending_vectorization_doc = None # Clear any previous pending state
                st.rerun() # Trigger a rerun to show the upload form in the sidebar

    db.close() # Close the session AFTER the loop has finished accessing relationships

# The "Chat" page placeholder (if you keep it)
if page == "Chat":
    st.title("Chat")
    st.info("This is the Chat page.")